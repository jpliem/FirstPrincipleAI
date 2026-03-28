from pathlib import Path

import mistune
from mistune.plugins.table import table as table_plugin
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def render_markdown_to_docx(
    markdown_text: str,
    output_path: str,
    template_path: str | None = None,
    title: str | None = None,
    metadata: dict | None = None,
):
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()

    # mistune v3: use Markdown.parse() to get AST tokens
    md = mistune.Markdown()
    table_plugin(md)
    tokens, _state = md.parse(markdown_text)

    if metadata:
        meta_para = doc.add_paragraph()
        meta_para.style = doc.styles["Normal"]
        for key, value in metadata.items():
            run = meta_para.add_run(f"{key}: {value}\n")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

    for token in tokens:
        _render_token(doc, token)

    doc.save(output_path)


def _render_token(doc: Document, token: dict):
    token_type = token.get("type", "")

    if token_type == "heading":
        level = token.get("attrs", {}).get("level", 1)
        text = _extract_text(token.get("children", []))
        doc.add_heading(text, level=min(level, 3))

    elif token_type == "paragraph":
        para = doc.add_paragraph()
        _render_inline(para, token.get("children", []))

    elif token_type == "list":
        ordered = token.get("attrs", {}).get("ordered", False)
        _render_list(doc, token, ordered)

    elif token_type == "block_text":
        # block_text wraps content inside some list items
        for child in token.get("children", []):
            _render_token(doc, child)

    elif token_type == "table":
        _render_table(doc, token)

    elif token_type == "block_code":
        code = token.get("raw", token.get("text", ""))
        para = doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        shading = run._element.get_or_add_rPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:fill"): "F0F0F0",
        })
        shading.append(shd)

    elif token_type == "thematic_break":
        doc.add_paragraph("─" * 50)


def _render_list(doc: Document, token: dict, ordered: bool, level: int = 0):
    """Render a list, handling nested lists and various child node types."""
    for item in token.get("children", []):
        if item.get("type") != "list_item":
            continue

        para = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        # Indent nested lists
        if level > 0:
            para.paragraph_format.left_indent = Pt(18 * level)

        children = item.get("children", [])
        for child in children:
            child_type = child.get("type", "")

            if child_type == "paragraph":
                _render_inline(para, child.get("children", []))

            elif child_type == "block_text":
                # block_text wraps inline content in some list contexts
                for block_child in child.get("children", []):
                    if block_child.get("type") == "paragraph":
                        _render_inline(para, block_child.get("children", []))
                    elif block_child.get("type") == "list":
                        nested_ordered = block_child.get("attrs", {}).get("ordered", False)
                        _render_list(doc, block_child, nested_ordered, level + 1)
                    else:
                        # Fallback: extract any text
                        text = _extract_text(block_child.get("children", []))
                        if not text:
                            text = block_child.get("raw", block_child.get("text", ""))
                        if text:
                            para.add_run(text)

            elif child_type == "list":
                # Nested list inside a list item
                nested_ordered = child.get("attrs", {}).get("ordered", False)
                _render_list(doc, child, nested_ordered, level + 1)

            else:
                # Fallback: try to extract text from unknown child types
                text = _extract_text(child.get("children", []))
                if not text:
                    text = child.get("raw", child.get("text", ""))
                if text:
                    para.add_run(text)


def _render_inline(para, children: list):
    for child in children:
        child_type = child.get("type", "")

        if child_type == "text":
            para.add_run(child.get("raw", child.get("text", "")))

        elif child_type == "strong":
            text = _extract_text(child.get("children", []))
            run = para.add_run(text)
            run.bold = True

        elif child_type == "emphasis":
            text = _extract_text(child.get("children", []))
            run = para.add_run(text)
            run.italic = True

        elif child_type == "codespan":
            text = child.get("raw", child.get("text", ""))
            run = para.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif child_type == "link":
            # Render link text, and URL in parentheses if present
            text = _extract_text(child.get("children", []))
            url = child.get("attrs", {}).get("url", child.get("link", ""))
            para.add_run(text)
            if url and url != text:
                run = para.add_run(f" ({url})")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(100, 100, 200)

        elif child_type == "image":
            alt = child.get("attrs", {}).get("alt", child.get("alt", ""))
            para.add_run(f"[Image: {alt}]" if alt else "[Image]")

        elif child_type == "softbreak":
            para.add_run("\n")

        elif child_type == "linebreak":
            para.add_run("\n")

        else:
            # Fallback — extract any text from unknown inline types
            text = _extract_text(child.get("children", []))
            if not text:
                text = child.get("raw", child.get("text", ""))
            if text:
                para.add_run(text)


def _extract_text(children: list) -> str:
    parts = []
    for child in children:
        if child.get("type") == "text":
            parts.append(child.get("raw", child.get("text", "")))
        elif "children" in child:
            parts.append(_extract_text(child["children"]))
        elif "raw" in child:
            parts.append(child["raw"])
    return "".join(parts)


def _render_table(doc: Document, token: dict):
    children = token.get("children", [])
    if not children:
        return

    rows = []
    for child in children:
        child_type = child.get("type", "")
        if child_type == "table_head":
            # In mistune v3, table_head contains table_cell directly (no table_row wrapper)
            cells = []
            for cell in child.get("children", []):
                if cell.get("type") == "table_cell":
                    cells.append(_extract_text(cell.get("children", [])))
            if cells:
                rows.append(cells)
        elif child_type == "table_body":
            for row in child.get("children", []):
                if row.get("type") == "table_row":
                    cells = []
                    for cell in row.get("children", []):
                        if cell.get("type") == "table_cell":
                            cells.append(_extract_text(cell.get("children", [])))
                    rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
