from pathlib import Path

import httpx

from app.config import settings


def parse_document(file_path: str, content_type: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md") or content_type in ("text/plain", "text/markdown"):
        return path.read_text(encoding="utf-8")

    if suffix == ".csv" or content_type == "text/csv":
        return path.read_text(encoding="utf-8")

    if suffix == ".pdf" or content_type == "application/pdf":
        return _parse_pdf(file_path)

    if suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(file_path)

    if suffix in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        return _parse_image_ocr(file_path)

    raise ValueError(f"Unsupported file type: {suffix} ({content_type})")


def _parse_pdf(file_path: str) -> str:
    import fitz

    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            text_parts.append(text)

    full_text = "\n\n".join(text_parts)

    if len(full_text.strip()) < 100 and settings.ocr_service_url:
        return _parse_image_ocr(file_path)

    return full_text


def _parse_docx(file_path: str) -> str:
    import docx

    doc = docx.Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _parse_image_ocr(file_path: str) -> str:
    if not settings.ocr_service_url:
        raise ValueError("OCR service URL not configured. Cannot process image/scanned documents.")

    with open(file_path, "rb") as f:
        response = httpx.post(
            settings.ocr_service_url,
            files={"file": (Path(file_path).name, f)},
            timeout=120.0,
        )
    response.raise_for_status()
    return response.json().get("text", "")
