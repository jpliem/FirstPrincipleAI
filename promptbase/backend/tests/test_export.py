import tempfile
from pathlib import Path

import docx

from app.export.renderer import render_markdown_to_docx


def test_render_headings():
    md = "# Title\n\n## Subtitle\n\n### Section\n\nParagraph text."
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        render_markdown_to_docx(md, f.name)
        path = Path(f.name)
    assert path.exists()
    assert path.stat().st_size > 0

    doc = docx.Document(f.name)
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert "Heading 2" in styles


def test_render_table():
    md = "| Name | Value |\n|------|-------|\n| Foo | 1 |\n| Bar | 2 |"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        render_markdown_to_docx(md, f.name)

    doc = docx.Document(f.name)
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "Name"


def test_render_bold_italic():
    md = "This is **bold** and *italic* text."
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        render_markdown_to_docx(md, f.name)

    doc = docx.Document(f.name)
    runs = doc.paragraphs[0].runs
    bold_found = any(r.bold for r in runs)
    italic_found = any(r.italic for r in runs)
    assert bold_found
    assert italic_found


def test_render_bullet_list():
    md = "- Item one\n- Item two\n- Item three"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        render_markdown_to_docx(md, f.name)

    doc = docx.Document(f.name)
    list_items = [p for p in doc.paragraphs if "List" in (p.style.name or "")]
    assert len(list_items) >= 3


def test_render_code_block():
    md = "```python\ndef hello():\n    print('hi')\n```"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        render_markdown_to_docx(md, f.name)

    doc = docx.Document(f.name)
    found_code = any("hello" in p.text for p in doc.paragraphs)
    assert found_code


def test_render_with_template():
    md = "# Test\n\nContent here."
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        render_markdown_to_docx(md, f.name, template_path=None)
    assert Path(f.name).stat().st_size > 0
